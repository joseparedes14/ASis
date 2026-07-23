"""
Document content extraction service.

Extracts text content from various file formats:
- PDF (pymupdf)
- DOCX (python-docx)
- TXT/MD (direct read)
- CSV (pandas)
- Excel (pandas + openpyxl)
- Images (rapidocr-onnxruntime for OCR)
"""

from pathlib import Path
from typing import Optional

from app.config.logging_config import get_logger

logger = get_logger(__name__)

# Extension to extractor mapping
EXTRACTORS: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "text",
    ".md": "text",
    ".csv": "csv",
    ".xlsx": "excel",
    ".xls": "excel",
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
}


class DocumentExtractor:
    """Extracts text content from various document formats.

    Handles PDF, DOCX, TXT, Markdown, CSV, Excel, and images (OCR).
    Returns extracted text along with metadata for classification.
    """

    def __init__(self) -> None:
        self._ocr_engine = None

    def _get_ocr_engine(self):
        """Lazy-load the OCR engine to avoid startup overhead."""
        if self._ocr_engine is None:
            try:
                from rapidocr_onnxruntime import RapidOCR
                self._ocr_engine = RapidOCR()
                logger.info("OCR engine initialized (RapidOCR)")
            except ImportError:
                logger.warning("RapidOCR not available, OCR will return empty text")
                return None
        return self._ocr_engine

    def extract(self, file_path: Path) -> Optional[str]:
        """Extract text content from a file.

        Args:
            file_path: Path to the file to extract content from.

        Returns:
            Extracted text content, or None if extraction failed.
        """
        if not file_path.exists():
            logger.warning("File does not exist: %s", file_path)
            return None

        suffix = file_path.suffix.lower()
        extractor_type = EXTRACTORS.get(suffix)

        if not extractor_type:
            logger.warning("Unsupported file type: %s", suffix)
            return None

        try:
            if extractor_type == "pdf":
                return self._extract_pdf(file_path)
            elif extractor_type == "docx":
                return self._extract_docx(file_path)
            elif extractor_type == "text":
                return self._extract_text(file_path)
            elif extractor_type == "csv":
                return self._extract_csv(file_path)
            elif extractor_type == "excel":
                return self._extract_excel(file_path)
            elif extractor_type == "image":
                return self._extract_image(file_path)
        except Exception as e:
            logger.error("Extraction failed for %s: %s", file_path, e)
            return None

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using pymupdf."""
        import pymupdf

        doc = pymupdf.open(str(file_path))
        text_parts = []

        for i, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"[Página {i + 1}]\n{page_text}")

        doc.close()
        return "\n\n".join(text_parts) if text_parts else ""

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document

        doc = Document(str(file_path))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        return file_path.read_text(encoding="utf-8", errors="replace")

    def _extract_csv(self, file_path: Path) -> str:
        """Extract content from CSV using pandas."""
        import pandas as pd

        df = pd.read_csv(str(file_path), nrows=50)  # Limit rows for context
        return df.to_string(index=False)

    def _extract_excel(self, file_path: Path) -> str:
        """Extract content from Excel using pandas."""
        import pandas as pd

        # Read first sheet, limit rows
        df = pd.read_excel(str(file_path), engine="openpyxl", nrows=50)
        return df.to_string(index=False)

    def _extract_image(self, file_path: Path) -> str:
        """Extract text from image using OCR (RapidOCR)."""
        ocr = self._get_ocr_engine()
        if ocr is None:
            return f"[Imagen: {file_path.name} — OCR no disponible]"

        result, elapse = ocr(str(file_path))
        if result is None:
            return f"[Imagen: {file_path.name} — No se detectó texto]"

        text_lines = [line[1] for line in result]
        return "\n".join(text_lines)
